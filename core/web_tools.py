# core/web_tools.py
import asyncio
import aiohttp
import requests
from bs4 import BeautifulSoup
import PyPDF2
import docx
from urllib.parse import urlparse, urljoin
import json
import os
from typing import Dict, List, Optional
import logging

class WebTools:
    def __init__(self):
        self.logger = logging.getLogger("WebTools")
        self.session = None
        
    async def __aenter__(self):
        self.session = aiohttp.ClientSession()
        return self
        
    async def __aexit__(self, exc_type, exc_val, exc_tb):
        if self.session:
            await self.session.close()

    async def web_search(self, query: str, num_results: int = 10) -> List[Dict]:
        """Perform web search using multiple search engines"""
        results = []
        
        # Try DuckDuckGo first (no API key required)
        try:
            ddg_results = await self.duckduckgo_search(query, num_results)
            results.extend(ddg_results)
        except Exception as e:
            self.logger.error(f"DuckDuckGo search failed: {e}")
        
        # Try Tavily if API key available
        if os.getenv("TAVILY_API_KEY"):
            try:
                tavily_results = await self.tavily_search(query, num_results)
                results.extend(tavily_results)
            except Exception as e:
                self.logger.error(f"Tavily search failed: {e}")
        
        # Try SerpAPI if available
        if os.getenv("SERPAPI_API_KEY"):
            try:
                serp_results = await self.serpapi_search(query, num_results)
                results.extend(serp_results)
            except Exception as e:
                self.logger.error(f"SerpAPI search failed: {e}")
        
        # Remove duplicates and return top results
        unique_results = []
        seen_urls = set()
        
        for result in results:
            if result.get("url") not in seen_urls:
                unique_results.append(result)
                seen_urls.add(result.get("url"))
                
        return unique_results[:num_results]

    async def duckduckgo_search(self, query: str, num_results: int = 10) -> List[Dict]:
        """Search using DuckDuckGo"""
        try:
            from duckduckgo_search import DDGS
            
            results = []
            with DDGS() as ddgs:
                for r in ddgs.text(query, max_results=num_results):
                    results.append({
                        "title": r.get("title", ""),
                        "url": r.get("href", ""),
                        "snippet": r.get("body", ""),
                        "source": "DuckDuckGo"
                    })
            return results
            
        except ImportError:
            self.logger.warning("duckduckgo_search not installed")
            return []
        except Exception as e:
            self.logger.error(f"DuckDuckGo search error: {e}")
            return []

    async def tavily_search(self, query: str, num_results: int = 10) -> List[Dict]:
        """Search using Tavily API"""
        try:
            url = "https://api.tavily.com/search"
            payload = {
                "api_key": os.getenv("TAVILY_API_KEY"),
                "query": query,
                "max_results": num_results,
                "search_depth": "advanced",
                "include_answer": True,
                "include_images": False,
                "include_raw_content": False
            }
            
            async with self.session.post(url, json=payload) as response:
                data = await response.json()
                
                results = []
                for item in data.get("results", []):
                    results.append({
                        "title": item.get("title", ""),
                        "url": item.get("url", ""),
                        "snippet": item.get("content", ""),
                        "source": "Tavily"
                    })
                
                return results
                
        except Exception as e:
            self.logger.error(f"Tavily search error: {e}")
            return []

    async def serpapi_search(self, query: str, num_results: int = 10) -> List[Dict]:
        """Search using SerpAPI (Google)"""
        try:
            url = "https://serpapi.com/search"
            params = {
                "api_key": os.getenv("SERPAPI_API_KEY"),
                "engine": "google",
                "q": query,
                "num": num_results
            }
            
            async with self.session.get(url, params=params) as response:
                data = await response.json()
                
                results = []
                for item in data.get("organic_results", []):
                    results.append({
                        "title": item.get("title", ""),
                        "url": item.get("link", ""),
                        "snippet": item.get("snippet", ""),
                        "source": "Google"
                    })
                
                return results
                
        except Exception as e:
            self.logger.error(f"SerpAPI search error: {e}")
            return []

    async def fetch_url_content(self, url: str) -> Dict:
        """Fetch and parse content from URL"""
        try:
            async with self.session.get(url, timeout=10) as response:
                content_type = response.headers.get('content-type', '').lower()
                
                if 'text/html' in content_type:
                    html = await response.text()
                    return self.parse_html_content(html, url)
                    
                elif 'application/pdf' in content_type:
                    pdf_content = await response.read()
                    return self.parse_pdf_content(pdf_content, url)
                    
                elif 'text/plain' in content_type:
                    text = await response.text()
                    return {
                        "url": url,
                        "title": urlparse(url).netloc,
                        "content": text,
                        "type": "text"
                    }
                    
                else:
                    return {
                        "url": url,
                        "title": urlparse(url).netloc,
                        "content": "Unsupported content type",
                        "type": "unknown"
                    }
                    
        except Exception as e:
            self.logger.error(f"Error fetching {url}: {e}")
            return {
                "url": url,
                "title": "Error",
                "content": f"Failed to fetch content: {str(e)}",
                "type": "error"
            }

    def parse_html_content(self, html: str, url: str) -> Dict:
        """Parse HTML content and extract text"""
        try:
            soup = BeautifulSoup(html, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get title
            title = soup.find('title')
            title_text = title.get_text().strip() if title else urlparse(url).netloc
            
            # Get main content
            # Try to find main content areas
            main_content = soup.find('main') or soup.find('article') or soup.find('div', class_='content')
            
            if main_content:
                content = main_content.get_text()
            else:
                content = soup.get_text()
            
            # Clean up text
            lines = (line.strip() for line in content.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            content = ' '.join(chunk for chunk in chunks if chunk)
            
            return {
                "url": url,
                "title": title_text,
                "content": content[:5000],  # Limit content length
                "type": "html"
            }
            
        except Exception as e:
            self.logger.error(f"Error parsing HTML from {url}: {e}")
            return {
                "url": url,
                "title": "Parse Error",
                "content": f"Failed to parse HTML: {str(e)}",
                "type": "error"
            }

    def parse_pdf_content(self, pdf_content: bytes, url: str) -> Dict:
        """Parse PDF content and extract text"""
        try:
            import io
            pdf_file = io.BytesIO(pdf_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return {
                "url": url,
                "title": f"PDF from {urlparse(url).netloc}",
                "content": text[:5000],  # Limit content length
                "type": "pdf"
            }
            
        except Exception as e:
            self.logger.error(f"Error parsing PDF from {url}: {e}")
            return {
                "url": url,
                "title": "PDF Parse Error",
                "content": f"Failed to parse PDF: {str(e)}",
                "type": "error"
            }

    async def summarize_content(self, content: str, max_length: int = 500) -> str:
        """Summarize content using extractive summarization"""
        try:
            from transformers import pipeline
            
            # Use a lightweight summarization model
            summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
            
            # Split content into chunks if too long
            max_chunk_length = 1024
            chunks = [content[i:i+max_chunk_length] for i in range(0, len(content), max_chunk_length)]
            
            summaries = []
            for chunk in chunks[:3]:  # Limit to first 3 chunks
                if len(chunk.strip()) > 50:  # Only summarize substantial chunks
                    summary = summarizer(chunk, max_length=max_length//len(chunks), min_length=30, do_sample=False)
                    summaries.append(summary[0]['summary_text'])
            
            return " ".join(summaries)
            
        except Exception as e:
            self.logger.error(f"Summarization error: {e}")
            # Fallback to simple truncation
            return content[:max_length] + "..." if len(content) > max_length else content

    async def process_file_upload(self, file_path: str) -> Dict:
        """Process uploaded file and extract content"""
        try:
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if file_extension == '.pdf':
                return self.process_pdf_file(file_path)
            elif file_extension in ['.docx', '.doc']:
                return self.process_word_file(file_path)
            elif file_extension == '.txt':
                return self.process_text_file(file_path)
            elif file_extension in ['.json']:
                return self.process_json_file(file_path)
            else:
                return {
                    "filename": os.path.basename(file_path),
                    "content": "Unsupported file type",
                    "type": "unsupported"
                }
                
        except Exception as e:
            self.logger.error(f"Error processing file {file_path}: {e}")
            return {
                "filename": os.path.basename(file_path),
                "content": f"Error processing file: {str(e)}",
                "type": "error"
            }

    def process_pdf_file(self, file_path: str) -> Dict:
        """Process PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                
                return {
                    "filename": os.path.basename(file_path),
                    "content": text,
                    "type": "pdf",
                    "pages": len(pdf_reader.pages)
                }
                
        except Exception as e:
            raise Exception(f"PDF processing error: {str(e)}")

    def process_word_file(self, file_path: str) -> Dict:
        """Process Word document"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return {
                "filename": os.path.basename(file_path),
                "content": text,
                "type": "docx",
                "paragraphs": len(doc.paragraphs)
            }
            
        except Exception as e:
            raise Exception(f"Word document processing error: {str(e)}")

    def process_text_file(self, file_path: str) -> Dict:
        """Process text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                
                return {
                    "filename": os.path.basename(file_path),
                    "content": content,
                    "type": "text",
                    "size": len(content)
                }
                
        except Exception as e:
            raise Exception(f"Text file processing error: {str(e)}")

    def process_json_file(self, file_path: str) -> Dict:
        """Process JSON file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
                
                # Convert JSON to readable text
                content = json.dumps(data, indent=2)
                
                return {
                    "filename": os.path.basename(file_path),
                    "content": content,
                    "type": "json",
                    "structure": type(data).__name__
                }
                
        except Exception as e:
            raise Exception(f"JSON file processing error: {str(e)}")

    async def get_latest_updates(self, topics: List[str]) -> List[Dict]:
        """Get latest updates on specified topics"""
        all_results = []
        
        for topic in topics:
            query = f"{topic} latest news updates"
            results = await self.web_search(query, num_results=5)
            
            for result in results:
                result["topic"] = topic
                all_results.append(result)
        
        return all_results

    async def research_topic(self, topic: str, depth: str = "basic") -> Dict:
        """Research a topic comprehensively"""
        
        # Basic search
        search_results = await self.web_search(f"{topic} comprehensive guide", num_results=10)
        
        # Fetch content from top results
        content_results = []
        for result in search_results[:5]:  # Limit to top 5 for performance
            content = await self.fetch_url_content(result["url"])
            content_results.append(content)
        
        # Combine all content
        combined_content = ""
        sources = []
        
        for content in content_results:
            if content["type"] != "error":
                combined_content += f"\n\n--- {content['title']} ---\n{content['content']}"
                sources.append({
                    "title": content["title"],
                    "url": content["url"],
                    "type": content["type"]
                })
        
        # Summarize if content is too long
        if len(combined_content) > 3000:
            summary = await self.summarize_content(combined_content, max_length=1500)
        else:
            summary = combined_content
        
        return {
            "topic": topic,
            "summary": summary,
            "sources": sources,
            "search_results": search_results
        }

# Example usage
async def main():
    async with WebTools() as web_tools:
        # Test web search
        results = await web_tools.web_search("Python machine learning", num_results=5)
        print("Search Results:")
        for result in results:
            print(f"- {result['title']}: {result['url']}")
        
        # Test content fetching
        if results:
            content = await web_tools.fetch_url_content(results[0]["url"])
            print(f"\nContent from {content['title']}:")
            print(content['content'][:200] + "...")
        
        # Test topic research
        research = await web_tools.research_topic("artificial intelligence")
        print(f"\nResearch Summary: {research['summary'][:200]}...")

if __name__ == "__main__":
    asyncio.run(main())